"""Post-process Pandoc DOCX for print-finished book interiors.

Applies book-design conventions Pandoc does not express well:

- Clear running heads and page numbers on cover / front matter
- Arabic page numbering restarting at the Introduction
- Suppress running heads on part/chapter/conclusion/appendix/bibliography
  openers (page numbers remain in the footer)
- Remove a Contents / TOC page when present (publication builds ship without a TOC)

Safe to run on any DOCX: missing landmarks make the corresponding step a no-op.
"""

from __future__ import annotations

import argparse
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

_BODY_OPENER_RE = re.compile(
    r"^(Introduction|Part\s+[IVXLCDM]+(?:\s+[—–-]\s+.+)?|Chapter\s+\d+|Conclusion|Appendix|"
    r"About the Series|Bibliography)$"
)


def _paragraph_style_name(paragraph: Paragraph) -> str:
    try:
        return paragraph.style.name if paragraph.style is not None else ""
    except Exception:
        return ""


def _paragraph_style_id(paragraph: Paragraph) -> str:
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return ""
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        return ""
    return str(p_style.get(qn("w:val")) or "")


def _is_body_opener(paragraph: Paragraph) -> bool:
    if _paragraph_style_name(paragraph) != "Heading 1":
        return False
    text = (paragraph.text or "").strip()
    return bool(_BODY_OPENER_RE.match(text))


def _find_intro_paragraph(doc: Document) -> Paragraph | None:
    for p in doc.paragraphs:
        if _paragraph_style_name(p) == "Heading 1" and (p.text or "").strip() == "Introduction":
            return p
    return None


def _find_contents_heading(doc: Document) -> Paragraph | None:
    for p in doc.paragraphs:
        if (p.text or "").strip() == "Contents" and _paragraph_style_name(p) in {
            "Heading 1",
            "Heading 2",
            "Title",
        }:
            return p
    for p in doc.paragraphs:
        if (p.text or "").strip() == "Contents":
            return p
    return None


def _is_page_break_only_paragraph(p_el) -> bool:
    """True when paragraph is only a page break (Pandoc ``\\newpage`` OpenXML)."""
    texts = [t.text for t in p_el.findall(".//" + qn("w:t")) if t.text]
    if any(t.strip() for t in texts):
        return False
    brs = p_el.findall(".//" + qn("w:br"))
    if not brs:
        return False
    return any(br.get(qn("w:type")) == "page" for br in brs)


def _clear_paragraph_content(p_el) -> None:
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)


def _make_sect_pr(
    template: OxmlElement,
    *,
    next_page: bool = True,
    start: int | None = None,
    title_page: bool = True,
) -> OxmlElement:
    sectPr = deepcopy(template)
    # Mid-document section properties should not keep stale header/footer rIds;
    # python-docx re-attaches them when we configure sections after reload.
    for tag in ("w:headerReference", "w:footerReference", "w:titlePg", "w:type", "w:pgNumType"):
        for el in list(sectPr.findall(qn(tag))):
            sectPr.remove(el)

    if next_page:
        type_el = OxmlElement("w:type")
        type_el.set(qn("w:val"), "nextPage")
        sectPr.insert(0, type_el)

    if title_page:
        sectPr.append(OxmlElement("w:titlePg"))

    if start is not None:
        pg = OxmlElement("w:pgNumType")
        pg.set(qn("w:start"), str(start))
        pg.set(qn("w:fmt"), "decimal")
        sectPr.append(pg)

    return sectPr


def _attach_sect_pr(p_el, sectPr: OxmlElement) -> None:
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    old = pPr.find(qn("w:sectPr"))
    if old is not None:
        pPr.remove(old)
    pPr.append(sectPr)


def _insert_section_break_before(paragraph: Paragraph, sectPr: OxmlElement) -> None:
    """Place a next-page section break immediately before *paragraph*.

    If the previous paragraph is a Pandoc page-break sentinel, convert it into
    the section break so chapters do not gain an extra blank page.
    """
    el = paragraph._element
    prev = el.getprevious()
    while prev is not None and prev.tag != qn("w:p"):
        prev = prev.getprevious()

    if prev is not None and _is_page_break_only_paragraph(prev):
        _clear_paragraph_content(prev)
        _attach_sect_pr(prev, sectPr)
        return

    p_break = OxmlElement("w:p")
    _attach_sect_pr(p_break, sectPr)
    el.addprevious(p_break)


def _add_page_number_field(paragraph: Paragraph) -> None:
    run1 = paragraph.add_run()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    run1._r.append(fc)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    run3._r.append(fc2)


def _wipe_container(container) -> None:
    for p in list(container.paragraphs):
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
    container.add_paragraph()


def _set_running_header(header, running_title: str) -> None:
    header.is_linked_to_previous = False
    _wipe_container(header)
    # _wipe left one empty paragraph
    hp = header.paragraphs[0]
    hp.alignment = 2  # RIGHT
    run = hp.add_run(running_title)
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Palatino Linotype"


def _set_page_footer(footer) -> None:
    footer.is_linked_to_previous = False
    _wipe_container(footer)
    fp = footer.paragraphs[0]
    fp.alignment = 1  # CENTER
    _add_page_number_field(fp)
    for run in fp.runs:
        run.font.size = Pt(9)
        run.font.name = "Palatino Linotype"


def _set_empty_header_footer(section) -> None:
    section.different_first_page_header_footer = False
    for part in (section.header, section.footer):
        part.is_linked_to_previous = False
        _wipe_container(part)


def _configure_body_section(section, running_title: str, *, restart_at: int | None) -> None:
    section.different_first_page_header_footer = True
    _set_running_header(section.header, running_title)
    _set_page_footer(section.footer)

    # Opener page: no running head; keep bottom page number.
    fp_header = section.first_page_header
    fp_header.is_linked_to_previous = False
    _wipe_container(fp_header)
    _set_page_footer(section.first_page_footer)

    sectPr = section._sectPr
    titlePg = sectPr.find(qn("w:titlePg"))
    if titlePg is None:
        sectPr.append(OxmlElement("w:titlePg"))

    pg = sectPr.find(qn("w:pgNumType"))
    if restart_at is not None:
        if pg is None:
            pg = OxmlElement("w:pgNumType")
            sectPr.append(pg)
        pg.set(qn("w:start"), str(restart_at))
        pg.set(qn("w:fmt"), "decimal")
    elif pg is not None and qn("w:start") in pg.attrib:
        del pg.attrib[qn("w:start")]


def _remove_contents_page(doc: Document) -> bool:
    """Remove the Contents heading and any following front-matter TOC body.

    Publication DOCX builds ship without a table of contents. Returns True when
    a Contents landmark was found and removed (including orphaned TOC fields).
    """
    contents = _find_contents_heading(doc)
    if contents is None:
        # Still strip any leftover TOC fields if Contents heading already gone.
        return _strip_toc_fields(doc)

    intro = _find_intro_paragraph(doc)
    remove: list[Paragraph] = [contents]
    seen_contents = False
    for p in doc.paragraphs:
        if p._element is contents._element:
            seen_contents = True
            continue
        if not seen_contents:
            continue
        if intro is not None and p._element is intro._element:
            break
        if _is_body_opener(p) and (p.text or "").strip() == "Introduction":
            break
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            break
        remove.append(p)

    for p in remove:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)

    _strip_toc_fields(doc)
    return True


def _strip_toc_fields(doc: Document) -> bool:
    """Remove paragraphs that contain a Word TOC field instruction."""
    removed = False
    for p in list(doc.paragraphs):
        instrs = p._element.findall(".//" + qn("w:instrText"))
        if not any(i.text and "TOC" in i.text for i in instrs):
            continue
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
            removed = True
    # Also drop orphaned TOC-style paragraphs left between Contents and Introduction
    # after a previous interior-finish pass (cached TOC lines with tab + dash).
    return removed


def _split_sections_at_openers(doc: Document) -> int:
    """Insert section breaks so each body opener starts a Word section.

    Returns the number of body openers found (0 if Introduction is missing).
    """
    body = doc.element.body
    final_sectPr = body.find(qn("w:sectPr"))
    if final_sectPr is None:
        return 0

    openers = [p for p in doc.paragraphs if _is_body_opener(p)]
    if not openers:
        return 0

    # Front matter ends immediately before Introduction. About the Series (and any
    # other recognized openers) may appear earlier in front matter; keep those in
    # the front-matter section and start body sections at Introduction.
    intro = next((p for p in openers if (p.text or "").strip() == "Introduction"), None)
    if intro is None:
        return 0
    body_openers = openers[openers.index(intro) :]

    # Insert from last opener backward so element positions stay stable.
    # The section break BEFORE opener[i] defines properties for the previous section.
    # final_sectPr defines properties for the last opener's section.
    for opener in reversed(body_openers[1:]):
        # Previous section is a body opener section (continuous page numbers).
        sectPr = _make_sect_pr(final_sectPr, next_page=True, start=None, title_page=True)
        _insert_section_break_before(opener, sectPr)

    # Break before Introduction: front-matter section (no title page / no nums).
    fm = _make_sect_pr(final_sectPr, next_page=True, start=None, title_page=False)
    _insert_section_break_before(intro, fm)

    # Final section (last opener) keeps final_sectPr; mark title page for opener head.
    for tag in ("w:titlePg", "w:pgNumType"):
        for el in list(final_sectPr.findall(qn(tag))):
            final_sectPr.remove(el)
    final_sectPr.append(OxmlElement("w:titlePg"))

    return len(body_openers)


def _apply_document_metadata(
    doc: Document,
    *,
    title: str,
    subtitle: str,
    author: str,
    keywords: str = "",
) -> None:
    """Set core properties used by Word, PDF converters, and accessibility tooling."""
    props = doc.core_properties
    props.title = title
    props.subject = subtitle
    props.author = author
    props.language = "en-US"
    props.keywords = keywords.strip() or (
        "history; power; democracy; leadership; institutions; "
        "collective action; shared power; public philosophy"
    )
    props.category = "Nonfiction"


def _strip_front_matter_image_captions(doc: Document) -> int:
    """Remove printed cover captions before the body (Introduction).

    Pandoc may emit an ``Image Caption`` paragraph from markdown alt text. That
    text belongs only in the drawing accessibility ``descr``, not on the page.
    Also drops empty ``Captioned Figure`` wrappers that contain no drawing once
    a caption was the sole sibling content (cover drawings are kept).
    """
    intro = _find_intro_paragraph(doc)
    removed = 0
    for p in list(doc.paragraphs):
        if intro is not None and p._element is intro._element:
            break
        style = _paragraph_style_name(p)
        style_id = _paragraph_style_id(p)
        text = (p.text or "").strip()
        has_drawing = bool(p._element.findall(".//" + qn("w:drawing")))
        is_image_caption = style == "Image Caption" or style_id == "ImageCaption"
        is_empty_captioned_figure = (
            (style == "Captioned Figure" or style_id == "CaptionedFigure")
            and not has_drawing
            and not text
        )
        if is_image_caption or is_empty_captioned_figure:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
                removed += 1
    return removed


def _set_first_cover_image_descr(doc: Document, cover_alt: str) -> bool:
    """Set accessibility descr/title on the first floating/inline drawing."""
    alt = cover_alt.strip()
    if not alt:
        return False
    for doc_pr in doc.element.body.iter(qn("wp:docPr")):
        doc_pr.set("descr", alt)
        # Keep title empty; screen readers use descr for the long description.
        if not doc_pr.get("title"):
            doc_pr.set("title", "")
        # Mirror onto pic:cNvPr when present (some readers prefer it).
        parent = doc_pr.getparent()
        if parent is not None:
            for cnv in parent.iter(qn("pic:cNvPr")):
                cnv.set("descr", alt)
        return True
    return False


def finish_interior_docx(
    path: Path,
    *,
    running_title: str,
    subtitle: str = "",
    author: str = "",
    keywords: str = "",
    cover_alt: str = "",
) -> dict:
    """Apply interior finish conventions to *path* in place."""
    status = {
        "body_openers": 0,
        "sections": 0,
        "front_matter_cleared": False,
        "body_headers": False,
        "toc_field": False,
        "toc_removed": False,
        "metadata": False,
        "cover_captions_removed": 0,
        "cover_descr_set": False,
    }

    doc = Document(str(path))
    n_openers = _split_sections_at_openers(doc)
    status["body_openers"] = n_openers
    if n_openers == 0:
        return status

    status["toc_removed"] = _remove_contents_page(doc)
    status["toc_field"] = False
    status["cover_captions_removed"] = _strip_front_matter_image_captions(doc)
    status["cover_descr_set"] = _set_first_cover_image_descr(doc, cover_alt)
    if running_title or subtitle or author or keywords:
        _apply_document_metadata(
            doc,
            title=running_title or "Manuscript",
            subtitle=subtitle or "",
            author=author or "",
            keywords=keywords or "",
        )
        status["metadata"] = True
    doc.save(str(path))

    doc = Document(str(path))
    status["sections"] = len(doc.sections)
    if len(doc.sections) < 2:
        return status

    _set_empty_header_footer(doc.sections[0])
    status["front_matter_cleared"] = True

    for i, section in enumerate(doc.sections[1:], start=1):
        restart = 1 if i == 1 else None
        _configure_body_section(section, running_title, restart_at=restart)
    status["body_headers"] = True

    doc.save(str(path))
    return status


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument(
        "--running-title",
        default="The Economy We Don't Experience",
    )
    args = parser.parse_args(argv)
    print(finish_interior_docx(args.docx, running_title=args.running_title))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
