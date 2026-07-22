"""Tests for DOCX interior finish (headers, sections, Contents removal)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_interior_finish import finish_interior_docx


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _build_mini_book(path: Path) -> None:
    doc = Document()
    # Cover image stand-in
    doc.add_paragraph("COVER")
    _add_page_break(doc)
    doc.add_heading("The Economy We Don't Experience", level=1)
    doc.add_heading("Subtitle", level=2)
    _add_page_break(doc)
    doc.add_heading("Copyright", level=1)
    doc.add_paragraph("Copyright text")
    _add_page_break(doc)
    doc.add_heading("Contents", level=1)
    doc.add_paragraph("Front Matter")
    doc.add_paragraph("Introduction — The Chart and the Receipt")
    doc.add_paragraph("Back Matter")
    doc.add_paragraph("Bibliography")
    _add_page_break(doc)
    doc.add_heading("Introduction", level=1)
    doc.add_heading("The Chart and the Receipt", level=2)
    doc.add_paragraph("Intro body.")
    _add_page_break(doc)
    doc.add_heading("Part I", level=1)
    doc.add_heading("The Economy We Describe", level=2)
    doc.add_paragraph("Part body.")
    _add_page_break(doc)
    doc.add_heading("Chapter 1", level=1)
    doc.add_heading("What the Average Leaves Out", level=2)
    doc.add_paragraph("Chapter body.")
    _add_page_break(doc)
    doc.add_heading("Bibliography", level=1)
    doc.add_paragraph("Bartels, Larry M.")
    # Seed a running header like reference.docx
    header = doc.sections[0].header
    header.paragraphs[0].text = "The Economy We Don't Experience"
    doc.save(str(path))


def test_finish_interior_clears_front_matter_and_splits_openers(tmp_path: Path) -> None:
    out = tmp_path / "book.docx"
    _build_mini_book(out)

    status = finish_interior_docx(out, running_title="The Economy We Don't Experience")
    assert status["body_openers"] == 4  # Intro, Part I, Ch1, Bibliography
    assert status["sections"] == 5
    assert status["front_matter_cleared"] is True
    assert status["body_headers"] is True
    assert status["toc_removed"] is True
    assert status["toc_field"] is False

    doc = Document(str(out))
    front = doc.sections[0]
    assert front.different_first_page_header_footer is False
    assert (front.header.paragraphs[0].text or "").strip() == ""
    assert (front.footer.paragraphs[0].text or "").strip() == ""

    body = doc.sections[1]
    assert body.different_first_page_header_footer is True
    assert "Economy" in (body.header.paragraphs[0].text or "")
    assert (body.first_page_header.paragraphs[0].text or "").strip() == ""

    pg = body._sectPr.find(qn("w:pgNumType"))
    assert pg is not None
    assert pg.get(qn("w:start")) == "1"

    texts = [(p.text or "").strip() for p in doc.paragraphs]
    assert "Contents" not in texts
    assert "Front Matter" not in texts
    joined = "\n".join(texts)
    assert "right-click" not in joined.lower()
    assert "Update Field" not in joined
    assert any(t == "Introduction" for t in texts)

    found_toc = False
    for p in doc.paragraphs:
        for instr in p._element.findall(".//" + qn("w:instrText")):
            if instr.text and "TOC" in instr.text:
                found_toc = True
                break
        if found_toc:
            break
    assert not found_toc


def test_finish_interior_noop_without_introduction(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("Not a book", level=1)
    doc.add_paragraph("Hello")
    path = tmp_path / "empty.docx"
    doc.save(str(path))
    status = finish_interior_docx(path, running_title="X")
    assert status["body_openers"] == 0
